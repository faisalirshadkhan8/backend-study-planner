"""
Document Processing Module

Handles file upload, text extraction, and preprocessing
for various document formats (PDF, TXT, DOCX).
"""

import os
import json
import hashlib
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
import logging

from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

from .config import RAGConfig
from pypdf import PdfReader


logger = logging.getLogger(__name__)


@dataclass
class DocumentMetadata:
    """Metadata for a processed document."""
    document_id: str
    filename: str
    file_size: int
    file_type: str
    upload_time: datetime
    processed_time: Optional[datetime] = None
    text_length: int = 0
    page_count: int = 0
    chunk_count: int = 0
    status: str = "uploaded"  # uploaded, processing, processed, error


class DocumentProcessor:
    """Handles document upload and text extraction."""
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self._ensure_directories()
    
    def _ensure_directories(self) -> None:
        """Create necessary directories if they don't exist."""
        for directory in [
            self.config.upload_folder,
            self.config.processed_folder,
            self.config.metadata_folder
        ]:
            os.makedirs(directory, exist_ok=True)
    
    def _generate_document_id(self, filename: str, content: bytes) -> str:
        """Generate unique document ID based on filename and content hash."""
        content_hash = hashlib.md5(content).hexdigest()[:8]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = secure_filename(filename).rsplit('.', 1)[0]
        return f"{safe_filename}_{timestamp}_{content_hash}"
    
    def _validate_file(self, file: FileStorage) -> Tuple[bool, str]:
        """Validate uploaded file."""
        if not file or not file.filename:
            return False, "No file provided"
        
        # Check file extension
        filename = file.filename.lower()
        if not any(filename.endswith(ext) for ext in self.config.allowed_extensions):
            return False, f"File type not allowed. Supported: {', '.join(self.config.allowed_extensions)}"
        
        # Check file size (read content to get actual size)
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > self.config.max_file_size:
            return False, f"File too large. Max size: {self.config.max_file_size // (1024*1024)}MB"
        
        if file_size == 0:
            return False, "Empty file"
        
        return True, "Valid"
    
    def upload_document(self, file: FileStorage) -> Tuple[bool, str, Optional[DocumentMetadata]]:
        """
        Upload and process a document.
        
        Returns:
            (success, message, metadata)
        """
        try:
            # Validate file
            is_valid, message = self._validate_file(file)
            if not is_valid:
                return False, message, None
            
            # Read file content
            content = file.read()
            file.seek(0)  # Reset for potential re-reading
            
            # Generate document ID and metadata
            doc_id = self._generate_document_id(file.filename, content)
            
            metadata = DocumentMetadata(
                document_id=doc_id,
                filename=file.filename,
                file_size=len(content),
                file_type=file.filename.rsplit('.', 1)[1].lower(),
                upload_time=datetime.now()
            )
            
            # Save raw file
            raw_path = os.path.join(self.config.upload_folder, f"{doc_id}.{metadata.file_type}")
            with open(raw_path, 'wb') as f:
                f.write(content)
            
            # Extract text
            text_content = self._extract_text(raw_path, metadata.file_type)
            if not text_content.strip():
                return False, "No text content could be extracted", None
            
            # Save processed text
            processed_path = os.path.join(self.config.processed_folder, f"{doc_id}.txt")
            with open(processed_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            # Update metadata
            metadata.processed_time = datetime.now()
            metadata.text_length = len(text_content)
            metadata.status = "processed"
            
            # Save metadata
            self._save_metadata(metadata)
            
            logger.info(f"Successfully processed document: {doc_id}")
            return True, f"Document processed successfully. ID: {doc_id}", metadata
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            return False, f"Processing error: {str(e)}", None
    
    def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from various file formats."""
        try:
            if file_type == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_type == 'txt':
                return self._extract_txt_text(file_path)
            elif file_type == 'docx':
                return self._extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            raise
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        reader = PdfReader(file_path)
        texts = []
        for page in reader.pages:
            txt = page.extract_text() or ""
            texts.append(txt)
        raw_text = "\n\n".join(texts)
        
        # Clean up spaced-out characters (e.g., "A s s i g n m e n t" → "Assignment")
        cleaned_text = self._clean_spaced_text(raw_text)
        
        # Remove front matter (TOC, title pages, etc.) before returning
        return self._remove_front_matter(cleaned_text)
    
    @staticmethod
    def _clean_spaced_text(text: str) -> str:
        """
        Fix text where each character is separated by spaces.
        Example: "A s s i g n m e n t" → "Assignment"
        
        This handles PDFs with poor text extraction.
        """
        import re
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Check if line has the pattern of single chars separated by spaces
            # Pattern: single letter/number, space, repeat
            if re.match(r'^(\w\s){4,}', line):  # At least 4 spaced characters
                # Remove all single spaces between single characters
                cleaned = re.sub(r'(\w)\s+(?=\w\s|\w$)', r'\1', line)
                cleaned_lines.append(cleaned)
            else:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            raw_text = file.read()
        
        # Remove front matter from TXT files too
        return self._remove_front_matter(raw_text)
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file using python-docx.

        Captures paragraphs and table cell text. DOCX has no stable page concept,
        so page markers are not added.
        """
        try:
            from docx import Document  # type: ignore
        except Exception as e:
            raise NotImplementedError(
                "DOCX support requires python-docx. Install with: pip install python-docx"
            ) from e

        doc = Document(file_path)
        parts: List[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                parts.append(text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = [(cell.text or "").strip() for cell in row.cells]
                line = " | ".join([c for c in cells if c])
                if line:
                    parts.append(line)

        raw_text = "\n".join(parts)
        
        # Remove front matter from DOCX files too
        return self._remove_front_matter(raw_text)
    
    def _remove_front_matter(self, text: str) -> str:
        """
        Remove front matter (title pages, TOC, certificates, etc.) from document text.
        Keeps only the substantive content starting from Chapter 1 or Introduction.
        
        This is applied to ALL document types (PDF, TXT, DOCX) during extraction.
        """
        import re
        
        lines = text.split('\n')
        content_start_idx = 0
        
        # Strategy 1: Find "Chapter 1" or "1 INTRODUCTION" or "CHAPTER 1"
        for i, line in enumerate(lines):
            line_upper = line.strip().upper()
            # Match patterns like: "CHAPTER 1", "1 INTRODUCTION", "Chapter 1:", etc.
            if re.match(r'^(CHAPTER\s+1|1[\s.:]+INTRODUCTION|CHAPTER\s+ONE)', line_upper):
                content_start_idx = i
                logger.info(f"Found content start at line {i}: '{line[:50]}'")
                break
            # Also catch numbered sections like "1.0" or "1.1"
            if re.match(r'^1\.[0-9]', line.strip()) and len(line.strip()) > 5:
                content_start_idx = i
                logger.info(f"Found content start (section 1.x) at line {i}: '{line[:50]}'")
                break
        
        # Strategy 2: If no chapter found, skip obvious front matter sections
        if content_start_idx == 0:
            skip_keywords = [
                'TABLE OF CONTENTS', 'CONTENTS', 'LIST OF FIGURES', 'LIST OF TABLES',
                'PLAGIARISM', 'CERTIFICATE', 'DECLARATION', 'ACKNOWLEDGEMENT',
                'DEDICATION', 'APPROVAL', 'ABSTRACT'
            ]
            
            in_front_matter = True
            for i, line in enumerate(lines):
                line_upper = line.strip().upper()
                
                # Check if we're still in front matter section
                if any(keyword in line_upper for keyword in skip_keywords):
                    in_front_matter = True
                    continue
                
                # If we find substantial content (long paragraph), assume front matter ended
                if in_front_matter and len(line.strip()) > 100 and not re.search(r'\.(\s*\.){2,}', line):
                    # This looks like real content (long line, no TOC dots)
                    content_start_idx = max(0, i - 2)  # Include a bit before
                    logger.info(f"Detected content start via paragraph at line {i}")
                    break
        
        # Strategy 3: Skip TOC entries (lines with multiple dots and page numbers)
        if content_start_idx == 0:
            for i, line in enumerate(lines):
                # If we see 5+ consecutive lines without TOC patterns, assume content started
                if i > 20:  # Give at least 20 lines for front matter
                    toc_count = 0
                    for j in range(max(0, i-5), i):
                        if re.search(r'\.(\s*\.){2,}|\d+$', lines[j]):
                            toc_count += 1
                    
                    if toc_count == 0:  # No TOC patterns in last 5 lines
                        content_start_idx = max(0, i - 5)
                        logger.info(f"Detected content start via TOC absence at line {i}")
                        break
        
        # If we found a content start, trim everything before it
        if content_start_idx > 0:
            skipped_lines = content_start_idx
            content = '\n'.join(lines[content_start_idx:])
            logger.info(f"Removed {skipped_lines} lines of front matter. Remaining: {len(content)} chars")
            return content
        
        # Fallback: return original if we couldn't detect front matter
        logger.warning("Could not detect front matter boundaries, returning full text")
        return text
    
    def _save_metadata(self, metadata: DocumentMetadata) -> None:
        """Save document metadata to JSON file."""
        metadata_path = os.path.join(
            self.config.metadata_folder, 
            f"{metadata.document_id}.json"
        )
        
        # Convert dataclass to dict for JSON serialization
        metadata_dict = {
            'document_id': metadata.document_id,
            'filename': metadata.filename,
            'file_size': metadata.file_size,
            'file_type': metadata.file_type,
            'upload_time': metadata.upload_time.isoformat(),
            'processed_time': metadata.processed_time.isoformat() if metadata.processed_time else None,
            'text_length': metadata.text_length,
            'page_count': metadata.page_count,
            'chunk_count': metadata.chunk_count,
            'status': metadata.status
        }
        
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(metadata_dict, f, indent=2)
    
    def get_document_metadata(self, document_id: str) -> Optional[DocumentMetadata]:
        """Load document metadata by ID."""
        metadata_path = os.path.join(self.config.metadata_folder, f"{document_id}.json")
        
        if not os.path.exists(metadata_path):
            return None
        
        try:
            with open(metadata_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            return DocumentMetadata(
                document_id=data['document_id'],
                filename=data['filename'],
                file_size=data['file_size'],
                file_type=data['file_type'],
                upload_time=datetime.fromisoformat(data['upload_time']),
                processed_time=datetime.fromisoformat(data['processed_time']) if data['processed_time'] else None,
                text_length=data['text_length'],
                page_count=data['page_count'],
                chunk_count=data['chunk_count'],
                status=data['status']
            )
        except Exception as e:
            logger.error(f"Failed to load metadata for {document_id}: {str(e)}")
            return None
    
    def list_documents(self) -> List[DocumentMetadata]:
        """List all processed documents."""
        documents = []
        
        if not os.path.exists(self.config.metadata_folder):
            return documents
        
        for filename in os.listdir(self.config.metadata_folder):
            if filename.endswith('.json'):
                doc_id = filename[:-5]  # Remove .json extension
                metadata = self.get_document_metadata(doc_id)
                if metadata:
                    documents.append(metadata)
        
        # Sort by upload time (newest first)
        documents.sort(key=lambda x: x.upload_time, reverse=True)
        return documents
    
    def delete_document(self, document_id: str) -> Tuple[bool, str]:
        """Delete document and all associated files."""
        try:
            metadata = self.get_document_metadata(document_id)
            if not metadata:
                return False, "Document not found"
            
            # Delete files
            files_to_delete = [
                os.path.join(self.config.upload_folder, f"{document_id}.{metadata.file_type}"),
                os.path.join(self.config.processed_folder, f"{document_id}.txt"),
                os.path.join(self.config.metadata_folder, f"{document_id}.json")
            ]
            
            for file_path in files_to_delete:
                if os.path.exists(file_path):
                    os.remove(file_path)
            
            logger.info(f"Successfully deleted document: {document_id}")
            return True, "Document deleted successfully"
            
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {str(e)}")
            return False, f"Deletion error: {str(e)}"
